"""DocumentGeneratorAgent — turns a finished report into a downloadable file.

The deep-research agent produces markdown; this agent renders that same
markdown as .md, .txt, .docx or .pdf so the user can take the report out of the
app. It is deterministic (no LLM call): the report is already written, and
re-generating prose per format would risk changing the content the user just
read and cited.

Markdown is parsed once into neutral blocks (`utils/markdown_blocks.py`), then
each writer walks those blocks. Adding a format means adding one writer, not
another markdown parser.

Heavy, optional dependencies (python-docx, reportlab) are imported inside their
writers so a missing package degrades to a clear 503 for that one format
instead of breaking app startup.
"""

import io
import re
from datetime import datetime

from backend.utils.logger import logger
from backend.utils.markdown_blocks import Span, parse_markdown, spans_to_text

# format → (extension, MIME type, human label)
FORMATS = {
    "md": ("md", "text/markdown; charset=utf-8", "Markdown"),
    "txt": ("txt", "text/plain; charset=utf-8", "Plain text"),
    "docx": (
        "docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "Word document",
    ),
    "pdf": ("pdf", "application/pdf", "PDF"),
}

# Reports are model-generated prose; anything past this is a bug or an abuse
# attempt, not a report. Guards the (CPU-bound) docx/pdf writers.
MAX_MARKDOWN_CHARS = 400_000


class UnsupportedFormatError(ValueError):
    """Requested format is not one of FORMATS."""


class FormatUnavailableError(RuntimeError):
    """The format is supported but its writer's dependency isn't installed."""


def safe_filename(title: str, extension: str) -> str:
    """Slugify a report title into a download filename. ASCII-only: the header
    value must survive latin-1 encoding in the Content-Disposition header."""
    base = re.sub(r"[^\w\s-]", "", (title or "").strip(), flags=re.ASCII)
    base = re.sub(r"[\s_-]+", "-", base).strip("-").lower()[:60]
    return f"{base or 'research-report'}.{extension}"


class DocumentGeneratorAgent:
    """Renders report markdown into a document. Stateless — one instance is
    shared by every request."""

    def generate(self, markdown: str, title: str = "", sources=None, fmt: str = "md"):
        """Return (bytes, mime_type, filename) for the requested format.

        `sources` is the research agent's list of {title, url} dicts; they are
        appended as a References section because the app shows them beside the
        report rather than inside it, and an exported file has no such sidebar.
        """
        fmt = (fmt or "").lower().strip()
        if fmt not in FORMATS:
            raise UnsupportedFormatError(f"Unsupported export format: {fmt!r}")

        markdown = markdown or ""
        if len(markdown) > MAX_MARKDOWN_CHARS:
            raise ValueError("Report is too large to export.")

        # The report's own H1 wins over the caller's title (which is the plan
        # title, often worded differently) — otherwise the document would open
        # with the plan title and then repeat the H1 as its first heading.
        title = (self._title_from_markdown(markdown) or title or "Research report").strip()
        sources = self._clean_sources(sources)
        extension, mime, _label = FORMATS[fmt]

        writer = {
            "md": self._write_md,
            "txt": self._write_txt,
            "docx": self._write_docx,
            "pdf": self._write_pdf,
        }[fmt]

        data = writer(markdown, title, sources)
        logger.info(
            f"[Generator] Exported '{title}' as .{extension} ({len(data)} bytes, "
            f"{len(sources)} sources)"
        )
        return data, mime, safe_filename(title, extension)

    # ── Shared helpers ────────────────────────────────────────────────────

    @staticmethod
    def _title_from_markdown(markdown: str) -> str:
        """Fall back to the report's own H1 when the caller sends no title."""
        for line in (markdown or "").split("\n"):
            if line.strip().startswith("# "):
                return line.strip()[2:].strip()
        return ""

    @staticmethod
    def _clean_sources(sources) -> list[dict]:
        """Keep only http(s) sources, de-duplicated, order preserved — the [n]
        markers in the report body refer to this order."""
        cleaned, seen = [], set()
        for s in sources or []:
            if not isinstance(s, dict):
                continue
            url = (s.get("url") or "").strip()
            if not url.lower().startswith(("http://", "https://")) or url in seen:
                continue
            seen.add(url)
            cleaned.append({"title": (s.get("title") or url).strip(), "url": url})
        return cleaned

    @staticmethod
    def _body_blocks(markdown: str, title: str) -> list[dict]:
        """Parsed blocks minus a leading H1 that just repeats `title` — every
        writer renders the title itself, so keeping it would duplicate it."""
        blocks = parse_markdown(markdown)
        if blocks and blocks[0]["type"] == "heading" and blocks[0]["level"] == 1:
            if spans_to_text(blocks[0]["spans"]).strip().lower() == title.strip().lower():
                return blocks[1:]
        return blocks

    @staticmethod
    def _subtitle() -> str:
        return f"Generated by Project Lovelace · {datetime.now().strftime('%d %B %Y')}"

    # ── Markdown ──────────────────────────────────────────────────────────

    def _write_md(self, markdown: str, title: str, sources: list[dict]) -> bytes:
        """Markdown passes through as-is; only the References section is added."""
        out = markdown.strip()
        if not out.lstrip().startswith("#"):
            out = f"# {title}\n\n{out}"
        if sources:
            lines = [f"{i}. [{s['title']}]({s['url']})" for i, s in enumerate(sources, 1)]
            out += "\n\n## References\n\n" + "\n".join(lines)
        return (out + "\n").encode("utf-8")

    # ── Plain text ────────────────────────────────────────────────────────

    def _write_txt(self, markdown: str, title: str, sources: list[dict]) -> bytes:
        blocks = self._body_blocks(markdown, title)
        out: list[str] = [title, "=" * min(len(title), 78), self._subtitle(), ""]

        for b in blocks:
            kind = b["type"]
            if kind == "heading":
                text = spans_to_text(b["spans"])
                if b["level"] <= 2:
                    out += ["", text, "-" * min(len(text), 78), ""]
                else:
                    out += ["", text, ""]
            elif kind == "paragraph":
                out += [self._wrap(spans_to_text(b["spans"])), ""]
            elif kind == "list_item":
                indent = "    " * b["level"]
                marker = f"{b['number']}." if b["ordered"] else "-"
                out.append(self._wrap(
                    f"{marker} {spans_to_text(b['spans'])}",
                    indent=indent,
                    hanging=" " * (len(marker) + 1),
                ))
            elif kind == "quote":
                out += [self._wrap(spans_to_text(b["spans"]), indent="  | "), ""]
            elif kind == "code":
                out += [""] + ["    " + l for l in b["text"].split("\n")] + [""]
            elif kind == "math":
                out += ["", "    " + b["text"], ""]
            elif kind == "rule":
                out += ["", "-" * 78, ""]
            elif kind == "table":
                out += self._txt_table(b) + [""]

        if sources:
            out += ["", "REFERENCES", "-" * 10, ""]
            out += [f"[{i}] {s['title']}\n    {s['url']}" for i, s in enumerate(sources, 1)]

        text = "\n".join(out).strip() + "\n"
        return re.sub(r"\n{4,}", "\n\n\n", text).encode("utf-8")

    @staticmethod
    def _wrap(text: str, width: int = 78, indent: str = "", hanging: str = "") -> str:
        """Greedy word wrap. `hanging` is added to continuation lines so list
        text stays aligned under its own first character, not under the bullet."""
        words, lines, current = text.split(), [], indent
        prefix = indent + hanging
        for w in words:
            candidate = f"{current} {w}" if current.strip() else current + w
            if len(candidate) > width and current.strip():
                lines.append(current.rstrip())
                current = prefix + w
            else:
                current = candidate
        if current.strip():
            lines.append(current.rstrip())
        return "\n".join(lines) if lines else ""

    @staticmethod
    def _txt_table(block: dict) -> list[str]:
        """Render a table as fixed-width columns, capped so wide tables don't
        blow past the 78-column body text."""
        rows = [[spans_to_text(c) for c in block["header"]]]
        rows += [[spans_to_text(c) for c in r] for r in block["rows"]]
        widths = [
            min(max((len(r[i]) for r in rows if i < len(r)), default=0), 28)
            for i in range(len(rows[0]))
        ]

        def line(cells):
            return "  ".join(
                (cells[i] if i < len(cells) else "")[:widths[i]].ljust(widths[i])
                for i in range(len(widths))
            ).rstrip()

        return ["", line(rows[0]), "  ".join("-" * w for w in widths)] + [
            line(r) for r in rows[1:]
        ]

    # ── DOCX ──────────────────────────────────────────────────────────────

    def _write_docx(self, markdown: str, title: str, sources: list[dict]) -> bytes:
        try:
            import docx
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, RGBColor, Inches
        except ImportError as e:
            raise FormatUnavailableError(
                "Word export needs the 'python-docx' package (pip install python-docx)."
            ) from e

        doc = docx.Document()
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing = 1.15

        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sub = doc.add_paragraph(self._subtitle())
        sub_run = sub.runs[0]
        sub_run.font.size = Pt(9)
        sub_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

        for b in self._body_blocks(markdown, title):
            kind = b["type"]
            if kind == "heading":
                p = doc.add_heading("", level=min(max(b["level"], 1), 4))
                self._docx_spans(p, b["spans"], doc)
            elif kind == "paragraph":
                self._docx_spans(doc.add_paragraph(), b["spans"], doc)
            elif kind == "list_item":
                style = "List Number" if b["ordered"] else "List Bullet"
                try:
                    p = doc.add_paragraph(style=style)
                except KeyError:  # style missing from the default template
                    p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25 + 0.25 * b["level"])
                self._docx_spans(p, b["spans"], doc)
            elif kind == "quote":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                for run in self._docx_spans(p, b["spans"], doc):
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
            elif kind == "code":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_after = Pt(10)
                run = p.add_run(b["text"])
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
            elif kind == "math":
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(b["text"])
                run.italic = True
            elif kind == "rule":
                doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif kind == "table":
                self._docx_table(doc, b)

        if sources:
            doc.add_page_break()
            doc.add_heading("References", level=1)
            for i, s in enumerate(sources, 1):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.add_run(f"[{i}] ").bold = True
                p.add_run(f"{s['title']} — ")
                self._docx_hyperlink(p, s["url"], s["url"])

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    @staticmethod
    def _docx_spans(paragraph, spans: list[Span], doc):
        """Add spans to a paragraph as formatted runs; returns the runs so the
        caller can apply block-level styling on top."""
        from docx.shared import Pt, RGBColor

        runs = []
        for s in spans:
            if not s.text:
                continue
            if s.link:
                DocumentGeneratorAgent._docx_hyperlink(paragraph, s.text, s.link)
                continue
            run = paragraph.add_run(s.text)
            run.bold = s.bold
            run.italic = s.italic
            if s.code:
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xB4, 0x25, 0x4A)
            runs.append(run)
        return runs

    @staticmethod
    def _docx_hyperlink(paragraph, text: str, url: str):
        """python-docx has no hyperlink API, so the w:hyperlink element and its
        relationship are built by hand."""
        from docx.oxml.shared import OxmlElement, qn

        r_id = paragraph.part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        link = OxmlElement("w:hyperlink")
        link.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        props = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1A56DB")
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        props.append(color)
        props.append(underline)
        run.append(props)

        text_el = OxmlElement("w:t")
        text_el.text = text
        run.append(text_el)
        link.append(run)
        paragraph._p.append(link)

    @staticmethod
    def _docx_table(doc, block: dict):
        from docx.shared import Pt

        headers = block["header"]
        table = doc.add_table(rows=1, cols=len(headers))
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        for cell, spans in zip(table.rows[0].cells, headers):
            cell.text = ""
            for run in DocumentGeneratorAgent._docx_spans(cell.paragraphs[0], spans, doc):
                run.bold = True

        for row in block["rows"]:
            cells = table.add_row().cells
            for cell, spans in zip(cells, row):
                cell.text = ""
                DocumentGeneratorAgent._docx_spans(cell.paragraphs[0], spans, doc)

        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        run.font.size = Pt(9.5)
        doc.add_paragraph()

    # ── PDF ───────────────────────────────────────────────────────────────

    def _write_pdf(self, markdown: str, title: str, sources: list[dict]) -> bytes:
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.platypus import (
                HRFlowable, KeepTogether, PageBreak, Paragraph, Preformatted,
                SimpleDocTemplate, Spacer,
            )
        except ImportError as e:
            raise FormatUnavailableError(
                "PDF export needs the 'reportlab' package (pip install reportlab)."
            ) from e

        body_font, bold_font, italic_font, mono_font = self._pdf_fonts()
        ss = getSampleStyleSheet()

        def style(name, **kw):
            kw.setdefault("fontName", body_font)
            return ParagraphStyle(name, parent=ss["Normal"], **kw)

        s_title = style("LTitle", fontName=bold_font, fontSize=22, leading=27, spaceAfter=4)
        s_sub = style("LSub", fontSize=9, leading=12, textColor=colors.HexColor("#6B7280"),
                      spaceAfter=18)
        s_body = style("LBody", fontSize=10.5, leading=16, spaceAfter=9)
        s_h = {
            1: style("LH1", fontName=bold_font, fontSize=16, leading=20, spaceBefore=16,
                     spaceAfter=8, textColor=colors.HexColor("#111827")),
            2: style("LH2", fontName=bold_font, fontSize=13.5, leading=18, spaceBefore=14,
                     spaceAfter=6, textColor=colors.HexColor("#1F2937")),
            3: style("LH3", fontName=bold_font, fontSize=11.5, leading=15, spaceBefore=10,
                     spaceAfter=4, textColor=colors.HexColor("#374151")),
        }
        s_list = style("LList", fontSize=10.5, leading=15, spaceAfter=4)
        s_quote = style("LQuote", fontName=italic_font, fontSize=10.5, leading=15,
                        leftIndent=14, spaceAfter=9, textColor=colors.HexColor("#4B5563"))
        s_math = style("LMath", fontName=italic_font, fontSize=11, leading=16,
                       alignment=TA_CENTER, spaceBefore=6, spaceAfter=10)
        s_code = ParagraphStyle("LCode", parent=ss["Code"], fontName=mono_font, fontSize=8.5,
                                leading=11.5, leftIndent=10, textColor=colors.HexColor("#1F2937"))
        s_cell = style("LCell", fontSize=8.8, leading=12)
        s_cell_h = style("LCellH", fontName=bold_font, fontSize=8.8, leading=12)

        story = [Paragraph(self._pdf_escape(title), s_title),
                 Paragraph(self._pdf_escape(self._subtitle()), s_sub)]

        for b in self._body_blocks(markdown, title):
            kind = b["type"]
            markup = self._pdf_markup_spans(b.get("spans", []), mono_font)
            if kind == "heading":
                story.append(Paragraph(markup, s_h.get(min(b["level"], 3), s_h[3])))
            elif kind == "paragraph":
                story.append(Paragraph(markup, s_body))
            elif kind == "list_item":
                bullet = f"{b['number']}." if b["ordered"] else "•"
                indent = 14 + 14 * b["level"]
                story.append(Paragraph(
                    markup, ParagraphStyle(f"li{indent}", parent=s_list, leftIndent=indent,
                                           bulletIndent=indent - 12),
                    bulletText=bullet,
                ))
            elif kind == "quote":
                story.append(Paragraph(markup, s_quote))
            elif kind == "math":
                story.append(Paragraph(self._pdf_escape(b["text"]), s_math))
            elif kind == "code":
                # Preformatted keeps the code's own line breaks; long lines are
                # hard-wrapped because it will not wrap them itself.
                story += [Spacer(1, 4), Preformatted(self._pdf_wrap_code(b["text"]), s_code),
                          Spacer(1, 8)]
            elif kind == "rule":
                story += [Spacer(1, 6),
                          HRFlowable(width="100%", thickness=0.6,
                                     color=colors.HexColor("#D1D5DB")),
                          Spacer(1, 10)]
            elif kind == "table":
                story.append(self._pdf_table(b, s_cell, s_cell_h, mono_font))
                story.append(Spacer(1, 10))

        if sources:
            story.append(PageBreak())
            story.append(Paragraph("References", s_h[1]))
            for i, s in enumerate(sources, 1):
                safe_url = self._pdf_escape(s["url"])
                story.append(KeepTogether([
                    Paragraph(
                        f'<b>[{i}]</b> {self._pdf_escape(s["title"])}<br/>'
                        f'<font size="8.5" color="#1A56DB">'
                        f'<link href="{safe_url}">{safe_url}</link></font>',
                        ParagraphStyle("LRef", parent=s_body, spaceAfter=7, leading=13),
                    )
                ]))

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=2.2 * cm, rightMargin=2.2 * cm,
            topMargin=2.0 * cm, bottomMargin=2.0 * cm,
            title=title, author="Project Lovelace",
        )
        doc.build(story, onFirstPage=self._pdf_footer(body_font),
                  onLaterPages=self._pdf_footer(body_font))
        return buf.getvalue()

    @staticmethod
    def _pdf_fonts():
        """Register DejaVu (bundled with matplotlib) so Greek letters and the
        maths symbols from `latex_to_text` render. Falls back to the built-in
        Helvetica/Courier, which are Latin-1 only."""
        fallback = ("Helvetica", "Helvetica-Bold", "Helvetica-Oblique", "Courier")
        try:
            import os
            import matplotlib
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont

            font_dir = os.path.join(os.path.dirname(matplotlib.__file__),
                                    "mpl-data", "fonts", "ttf")
            faces = {
                "LV-Body": "DejaVuSans.ttf",
                "LV-Bold": "DejaVuSans-Bold.ttf",
                "LV-Italic": "DejaVuSans-Oblique.ttf",
                "LV-Mono": "DejaVuSansMono.ttf",
            }
            registered = set(pdfmetrics.getRegisteredFontNames())
            for name, filename in faces.items():
                if name in registered:
                    continue
                path = os.path.join(font_dir, filename)
                if not os.path.exists(path):
                    return fallback
                pdfmetrics.registerFont(TTFont(name, path))
            return ("LV-Body", "LV-Bold", "LV-Italic", "LV-Mono")
        except Exception as e:
            logger.warning(f"[Generator] Unicode PDF fonts unavailable, using Helvetica: {e}")
            return fallback

    @staticmethod
    def _pdf_escape(text: str) -> str:
        """Escape for reportlab's mini-HTML paragraph markup."""
        return (text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    @classmethod
    def _pdf_markup_spans(cls, spans: list[Span], mono_font: str) -> str:
        """Spans → reportlab paragraph markup (<b>, <i>, <font>, <link>)."""
        out = []
        for s in spans:
            if not s.text:
                continue
            chunk = cls._pdf_escape(s.text)
            if s.code:
                chunk = f'<font face="{mono_font}" size="9" color="#B4254A">{chunk}</font>'
            if s.bold:
                chunk = f"<b>{chunk}</b>"
            if s.italic:
                chunk = f"<i>{chunk}</i>"
            if s.link:
                chunk = (f'<link href="{cls._pdf_escape(s.link)}" color="#1A56DB">'
                         f"{chunk}</link>")
            out.append(chunk)
        return "".join(out) or "&nbsp;"

    @staticmethod
    def _pdf_wrap_code(text: str, width: int = 92) -> str:
        """Hard-wrap over-long code lines — Preformatted clips instead of wrapping."""
        lines = []
        for line in text.split("\n"):
            while len(line) > width:
                lines.append(line[:width])
                line = line[width:]
            lines.append(line)
        return "\n".join(lines)

    @classmethod
    def _pdf_table(cls, block, s_cell, s_cell_h, mono_font):
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, Table, TableStyle

        data = [[Paragraph(cls._pdf_markup_spans(c, mono_font), s_cell_h)
                 for c in block["header"]]]
        data += [[Paragraph(cls._pdf_markup_spans(c, mono_font), s_cell) for c in row]
                 for row in block["rows"]]

        # Equal columns across the text frame: cell Paragraphs wrap, so an even
        # split never overflows the page however lopsided the content is.
        usable = 21.0 * cm - 4.4 * cm
        col_width = usable / max(len(block["header"]), 1)
        table = Table(data, colWidths=[col_width] * len(block["header"]), repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F3F4F6")),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D1D5DB")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("RIGHTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#FAFAFA")]),
        ]))
        return table

    @staticmethod
    def _pdf_footer(body_font: str):
        """Page-number callback for SimpleDocTemplate.build()."""
        def draw(canvas, doc):
            canvas.saveState()
            canvas.setFont(body_font, 8)
            canvas.setFillColorRGB(0.42, 0.45, 0.50)
            canvas.drawCentredString(doc.pagesize[0] / 2, 1.1 * 28.35,
                                     str(canvas.getPageNumber()))
            canvas.restoreState()
        return draw
